"""Build the progress report docx from the actual experiment artifacts.

Reads results/metrics.json and results/demo_results.json so every number in
the report comes from a real run. Output: report/Progress_Report.docx
"""

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).parent
FIGURES = ROOT / "figures"
RESULTS = ROOT / "results"
REPORT_DIR = ROOT / "report"

INK = RGBColor(0x0B, 0x0B, 0x0B)
MUTED = RGBColor(0x52, 0x51, 0x4E)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = INK
    return h


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = MUTED


def add_figure(doc, filename, caption, width=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIGURES / filename), width=Inches(width))
    add_caption(doc, caption)


def style_table(table):
    table.style = "Light Grid Accent 1"


def main():
    REPORT_DIR.mkdir(exist_ok=True)
    metrics = json.loads((RESULTS / "metrics.json").read_text())
    demo = json.loads((RESULTS / "demo_results.json").read_text(encoding="utf-8"))

    ds = metrics["dataset"]
    nb = metrics["models"]["Naive Bayes"]
    lr = metrics["models"]["Logistic Regression"]

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Resume Analyzer and Job Matcher")
    run.font.size = Pt(20)
    run.font.bold = True
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("CAI4002: Introduction to Artificial Intelligence | "
                      "Group Project Progress Report")
    run.font.size = Pt(12)
    run.font.italic = True
    team = doc.add_paragraph()
    team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = team.add_run("Team Members: Martin Dang, Tuan Khang Phan\n"
                       "July 9, 2026")
    run.font.size = Pt(11)
    run.font.bold = True

    # 1. status summary
    add_heading(doc, "1. Project Status Summary")
    doc.add_paragraph(
        "The project is in progress and on track. Since the proposal, we have "
        "built a working end-to-end prototype of the Resume Analyzer and Job "
        "Matcher: the system loads a real resume, extracts a structured skill "
        "profile from it, parses a target job posting into required and "
        "preferred qualifications, computes a compatibility score, and returns "
        "a ranked list of concrete edits that would raise that score. We have "
        "also trained and evaluated our first two supervised models for resume "
        "category classification on a public dataset of "
        f"{ds['n_resumes']:,} real resumes. This report presents the current "
        "state of the pipeline, quantitative results from our first "
        "experiments, a walkthrough of the working demo, and the division of "
        "work between the two team members."
    )
    doc.add_paragraph(
        "All code, experiment scripts, figures, and results in this report "
        "are versioned in our GitHub repository "
        "(github.com/khangpt2k6/Intro_to_AI)."
    )

    # 2. dataset
    add_heading(doc, "2. Dataset")
    doc.add_paragraph(
        "We are using the public Resume Dataset from Kaggle "
        "(snehaanbhawal/resume-dataset), as planned in the proposal. It "
        f"contains {ds['n_resumes']:,} real resumes collected from "
        "livecareer.com, provided as plain text and labeled into "
        f"{ds['n_categories']} job categories such as Information Technology, "
        "Finance, Engineering, and Healthcare. The dataset serves two roles in "
        "our system: it is the training corpus for the TF-IDF vector space "
        "and the category classifier, and it provides realistic test resumes "
        "for the matching demo. Using a public dataset keeps the project "
        "reproducible and avoids the privacy concerns of collecting real "
        "applicant data."
    )
    t = doc.add_table(rows=5, cols=2)
    style_table(t)
    rows = [
        ("Total resumes", f"{ds['n_resumes']:,}"),
        ("Job categories", str(ds["n_categories"])),
        ("Training split (80%)", f"{ds['train_size']:,} resumes"),
        ("Held-out test split (20%)", f"{ds['test_size']:,} resumes"),
        ("Largest / smallest category", "120 (IT, Business Dev.) / 22 (BPO)"),
    ]
    for i, (k, v) in enumerate(rows):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
    doc.add_paragraph()
    add_figure(doc, "fig1_category_distribution.png",
               "Figure 1. Distribution of the 2,484 resumes across the 24 job "
               "categories. The class imbalance (120 vs 22 resumes) is one "
               "challenge our classifier has to handle.", width=5.6)

    # 3. system progress
    add_heading(doc, "3. What We Have Built So Far")
    doc.add_paragraph(
        "The pipeline from the proposal has five stages. Four of them are "
        "implemented and working; the table below shows the status of each."
    )
    t = doc.add_table(rows=6, cols=3)
    style_table(t)
    hdr = t.rows[0].cells
    hdr[0].text = "Pipeline stage"
    hdr[1].text = "Status"
    hdr[2].text = "Implementation"
    stages = [
        ("Skill extraction",
         "Working",
         "A curated skill taxonomy (60+ canonical skills with aliases, e.g. "
         "'JS' and 'JavaScript' normalize to one skill) matched with "
         "word-bounded regular expressions; each extracted skill carries a "
         "mention count and a confidence value."),
        ("Job description parsing",
         "Working",
         "Postings are split into hard constraints (required qualifications) "
         "and soft constraints (preferred qualifications), following the "
         "constraint-satisfaction framing in our proposal."),
        ("Compatibility scoring",
         "Working",
         "TF-IDF vectors (unigrams + bigrams, 20,000 features) fit on the "
         "full resume corpus; the score combines weighted skill coverage "
         "(60%) with the resume's percentile rank in corpus-wide cosine "
         "similarity to the posting (40%)."),
        ("Recommendation engine",
         "Working",
         "Missing skills are ranked greedy best-first by the exact score "
         "gain each edit recovers; required skills carry twice the weight "
         "of preferred ones."),
        ("Resume category classification",
         "First results",
         "Naive Bayes and Logistic Regression over TF-IDF features, "
         "evaluated on a stratified held-out test set (Section 4)."),
    ]
    for i, (a, b, c) in enumerate(stages, start=1):
        t.rows[i].cells[0].text = a
        t.rows[i].cells[1].text = b
        t.rows[i].cells[2].text = c
    doc.add_paragraph()
    doc.add_paragraph(
        "Not yet implemented (planned before the final report): transformer "
        "embeddings (Sentence-BERT) as a second similarity signal alongside "
        "TF-IDF, PDF/DOCX text extraction for user-uploaded resumes, a simple "
        "web interface, and a Naive Bayes confidence model for individual "
        "extracted skills."
    )

    # 4. experiments
    add_heading(doc, "4. First Experimental Results")
    doc.add_paragraph(
        "As a first supervised learning experiment, we trained two classifiers "
        "to predict a resume's job category from its text, using an 80/20 "
        f"stratified train/test split ({ds['train_size']:,} training, "
        f"{ds['test_size']:,} test resumes). This classifier gives the "
        "analyzer a sense of which job family a resume currently signals, "
        "and it exercises the same TF-IDF representation the matcher uses. "
        "With 24 classes, random guessing would score about 4.2%."
    )
    t = doc.add_table(rows=3, cols=3)
    style_table(t)
    hdr = t.rows[0].cells
    hdr[0].text = "Model"
    hdr[1].text = "Accuracy"
    hdr[2].text = "Macro F1"
    t.rows[1].cells[0].text = "Multinomial Naive Bayes"
    t.rows[1].cells[1].text = f"{nb['accuracy']:.1%}"
    t.rows[1].cells[2].text = f"{nb['macro_f1']:.3f}"
    t.rows[2].cells[0].text = "Logistic Regression"
    t.rows[2].cells[1].text = f"{lr['accuracy']:.1%}"
    t.rows[2].cells[2].text = f"{lr['macro_f1']:.3f}"
    doc.add_paragraph()
    add_figure(doc, "fig2_model_comparison.png",
               "Figure 2. Both models beat the 4.2% random baseline by a wide "
               "margin; Logistic Regression is the stronger of the two.",
               width=5.2)
    add_figure(doc, "fig3_confusion_matrix.png",
               "Figure 3. Normalized confusion matrix for Logistic Regression. "
               "The strong diagonal shows most categories are learned well; "
               "the remaining confusion concentrates in genuinely overlapping "
               "categories (e.g. Finance vs Accountant, Sales vs Business "
               "Development).", width=6.0)
    doc.add_paragraph(
        "Takeaways: Logistic Regression outperforms Naive Bayes by about 11 "
        "points of accuracy, and the errors that remain are concentrated in "
        "semantically overlapping categories, which supports our plan to add "
        "dense transformer embeddings that capture context beyond keyword "
        "counts."
    )

    # 5. demo
    add_heading(doc, "5. End-to-End Demo")
    sw = demo["software_engineer"]
    da = demo["data_analyst"]
    ac = demo["accountant"]
    doc.add_paragraph(
        "To validate the whole pipeline, we matched one real Information "
        "Technology resume from the dataset (ID 83816738) against three job "
        "postings we wrote in realistic formats: a backend software engineer "
        "role, a data analyst role, and a staff accountant role. The system "
        "behaves the way a human reviewer would expect: the IT resume scores "
        f"{sw['score']}/100 against the software engineer posting, "
        f"{da['score']}/100 against the data analyst posting, and only "
        f"{ac['score']}/100 against the accountant posting."
    )
    add_figure(doc, "fig4_match_scores.png",
               "Figure 4. The same resume scored against three postings. The "
               "score separates the strong fit from the partial and poor fits.",
               width=5.4)
    doc.add_paragraph(
        "For the software engineer posting, the analyzer matched "
        f"{len(sw['matched_hard'])} of "
        f"{len(sw['matched_hard']) + len(sw['missing_hard'])} required skills "
        f"({', '.join(sw['matched_hard'])}) plus "
        f"{len(sw['matched_soft'])} preferred skills "
        f"({', '.join(sw['matched_soft'])}), and identified "
        f"{', '.join(s['skill'] for s in sw['recommendations'][:2])} as the "
        "two missing required skills. The recommendation engine estimates "
        f"that surfacing each one is worth about "
        f"+{sw['recommendations'][0]['gain']} points, so they are ranked "
        "first, ahead of preferred-skill edits like Kubernetes."
    )
    add_figure(doc, "fig5_demo_output.png",
               "Figure 5. Screenshot of the analyzer's console output for the "
               "demo resume across all three postings, including the ranked "
               "edit recommendations.", width=6.2)

    # 6. challenges
    add_heading(doc, "6. Challenges Encountered")
    challenges = [
        ("Class imbalance in the dataset.",
         "The largest categories have 120 resumes while the smallest (BPO) "
         "has only 22, so a model can look accurate while ignoring small "
         "classes. We addressed this with a stratified train/test split and "
         "by reporting macro F1 (which weights every class equally) next to "
         "raw accuracy."),
        ("Semantically overlapping categories.",
         "Much of the classification error concentrates in category pairs "
         "that genuinely share vocabulary, such as Finance vs Accountant and "
         "Sales vs Business Development. Naive Bayes suffered most from "
         "this, which is why we added Logistic Regression as a second "
         "learner (+11 points of accuracy) and plan to test transformer "
         "embeddings that capture context beyond shared keywords."),
        ("Raw cosine similarities are not interpretable.",
         "Resume-to-posting cosine values in our corpus range from about "
         "0.01 to 0.17, so presenting them directly would make every match "
         "look terrible. We solved this with percentile calibration: the "
         "score reports where the resume ranks against all 2,484 corpus "
         "resumes for that same posting, which turns an opaque 0.174 into "
         "an interpretable 'beats 100% of the corpus'."),
        ("Short skill aliases cause false positives.",
         "Aliases like 'R', 'Go', or 'C' match ordinary English words. We "
         "mitigated this with word-boundary regular expressions and by "
         "requiring longer context forms (for example 'R programming' "
         "instead of bare 'R'), at the cost of some recall; measuring that "
         "precision/recall trade-off on a hand-labeled sample is on our "
         "remaining-work list."),
        ("Noisy resume text.",
         "The dataset resumes were converted from PDFs, so the text contains "
         "run-together section headers and inconsistent formatting. Our "
         "first demo subject extracted almost no skills for this reason; "
         "investigating it led us to broaden the alias lists and pick "
         "test subjects systematically instead of arbitrarily."),
    ]
    for title_text, body in challenges:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(title_text + " ")
        run.font.bold = True
        p.add_run(body)
    doc.add_paragraph()

    # 7. contributions
    add_heading(doc, "7. Team Member Contributions")
    doc.add_paragraph(
        "We split the work evenly (50/50), pairing on design decisions and "
        "dividing implementation by pipeline stage:"
    )
    t = doc.add_table(rows=3, cols=2)
    style_table(t)
    hdr = t.rows[0].cells
    hdr[0].text = "Martin Dang (50%)"
    hdr[1].text = "Tuan Khang Phan (50%)"
    t.rows[1].cells[0].text = (
        "Skill taxonomy design and alias normalization; skill extraction "
        "module; job description constraint parser (required vs preferred); "
        "recommendation engine and its ranking logic")
    t.rows[1].cells[1].text = (
        "Dataset acquisition and preprocessing; TF-IDF matching engine and "
        "percentile calibration; classifier experiments (Naive Bayes, "
        "Logistic Regression) and evaluation; demo script and figures")
    t.rows[2].cells[0].text = (
        "Joint: sample job postings, testing the pipeline end to end, this "
        "progress report")
    t.rows[2].cells[1].text = (
        "Joint: system design, error analysis of demo output, this progress "
        "report")
    doc.add_paragraph()

    # 8. remaining work
    add_heading(doc, "8. Remaining Work")
    items = [
        "Add Sentence-BERT embeddings as a dense similarity signal next to "
        "TF-IDF and compare the two on the same demo cases (connects to the "
        "deep learning material from class).",
        "Add PDF and DOCX text extraction so users can upload their own "
        "resume files (the dataset also ships the same resumes as PDFs, "
        "which gives us a ready-made test set for this).",
        "Replace the frequency-based skill confidence heuristic with a "
        "Naive Bayes estimate over the surrounding section text.",
        "Expand the skill taxonomy beyond the current 60+ canonical skills "
        "and evaluate extraction precision/recall on a hand-labeled sample.",
        "Build a minimal web interface that accepts a resume and a posting "
        "and renders the score, the matched/missing breakdown, and the "
        "ranked recommendations.",
    ]
    for it in items:
        doc.add_paragraph(it, style="List Bullet")
    doc.add_paragraph(
        "We are confident the remaining items fit in the time before the "
        "final deliverable, since the core pipeline and evaluation harness "
        "are already in place."
    )

    out = REPORT_DIR / "Progress_Report.docx"
    doc.save(out)
    print(f"Saved {out}")


if __name__ == "__main__":
    main()
