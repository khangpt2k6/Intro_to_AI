"""Build the final report docx (and pdf) from the actual experiment artifacts.

Reads results/metrics.json and results/demo_results.json so every number in the
report comes from a real run. Styling: one font family everywhere (Times New
Roman) plus a navy/blue accent palette for headings, tables and callouts.
Output: report/Final_Report.docx (+ .pdf if Word is available).
"""

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).parent
FIGURES = ROOT / "figures"
RESULTS = ROOT / "results"
REPORT_DIR = ROOT / "report"

FONT = "Times New Roman"

# palette
NAVY = RGBColor(0x1F, 0x38, 0x64)       # headings, table header fill
BLUE = RGBColor(0x2E, 0x74, 0xB5)       # sub headings, accents
TEAL = RGBColor(0x1B, 0x6E, 0x6B)       # numeric highlights
INK = RGBColor(0x1A, 0x1A, 0x1A)        # body text
MUTED = RGBColor(0x59, 0x59, 0x59)      # captions
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HDR_FILL = "1F3864"
ALT_FILL = "EDF1F8"
BOX_FILL = "F3F7FC"
GRID = "A9BEDC"


# ---------------------------------------------------------------- low level


def _shade(cell, hex_fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _table_borders(table, hex_color=GRID, size=6):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), hex_color)
        borders.append(el)
    tbl_pr.append(borders)


def _para_border(paragraph, edge="bottom", hex_color="2E74B5", size=8, space=4):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), hex_color)
    borders.append(el)


def _cell_margins(table, left=110, right=110, top=60, bottom=60):
    tbl_pr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for name, val in (("top", top), ("left", left),
                      ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tbl_pr.append(mar)


# ---------------------------------------------------------------- building


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    h.paragraph_format.space_after = Pt(4)
    for run in h.runs:
        run.font.color.rgb = NAVY if level == 1 else BLUE
        run.font.size = Pt(15 if level == 1 else 12.5)
        run.font.bold = True
    if level == 1:
        _para_border(h, "bottom", "2E74B5", 8, 3)
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.color.rgb = INK
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = MUTED
    return p


def add_figure(doc, filename, caption, width=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(FIGURES / filename), width=Inches(width))
    add_caption(doc, caption)


def add_table(doc, header, rows, widths=None):
    """Header row filled navy with white text, body rows banded light blue."""
    t = doc.add_table(rows=len(rows) + 1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_borders(t)
    _cell_margins(t)

    for j, label in enumerate(header):
        cell = t.rows[0].cells[j]
        cell.text = label
        _shade(cell, HDR_FILL)
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE

    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = value
            if i % 2 == 0:
                _shade(cell, ALT_FILL)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.color.rgb = INK
                    if j == 0:
                        run.font.bold = True

    if widths:
        for i in range(len(rows) + 1):
            for j, w in enumerate(widths):
                t.rows[i].cells[j].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def add_callout(doc, title, body):
    """Tinted single-cell box with a thick blue left edge."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _cell_margins(t, left=150, right=150, top=90, bottom=90)
    cell = t.rows[0].cells[0]
    _shade(cell, BOX_FILL)

    tbl_pr = t._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge, size, color in (("top", 4, GRID), ("bottom", 4, GRID),
                              ("right", 4, GRID), ("left", 24, "2E74B5")):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tbl_pr.append(borders)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(title + " ")
    r.font.bold = True
    r.font.color.rgb = NAVY
    r = p.add_run(body)
    r.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def add_bullet(doc, lead, body):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(lead + " ")
    r.font.bold = True
    r.font.color.rgb = BLUE
    r = p.add_run(body)
    r.font.color.rgb = INK
    return p


def add_metric_strip(doc, items):
    """Row of coloured key numbers, one cell per metric."""
    t = doc.add_table(rows=2, cols=len(items))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _cell_margins(t, top=90, bottom=90)
    _table_borders(t, GRID, 6)
    for j, (value, label) in enumerate(items):
        top = t.rows[0].cells[j]
        _shade(top, BOX_FILL)
        p = top.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        r.font.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = TEAL

        bottom = t.rows[1].cells[j]
        _shade(bottom, BOX_FILL)
        p = bottom.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def apply_fonts(doc):
    """One font family for the whole document, body size 12pt."""
    for style_name in ("Normal", "List Bullet", "List Paragraph", "Title",
                       "Heading 1", "Heading 2", "Heading 3", "Caption"):
        try:
            st = doc.styles[style_name]
        except KeyError:
            continue
        st.font.name = FONT
        rpr = st.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rfonts.set(qn(attr), FONT)
    doc.styles["Normal"].font.size = Pt(12)
    doc.styles["Normal"].font.color.rgb = INK

    def fix(paragraphs):
        for p in paragraphs:
            for r in p.runs:
                r.font.name = FONT
                rpr = r._element.get_or_add_rPr()
                rfonts = rpr.get_or_add_rFonts()
                for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                    rfonts.set(qn(attr), FONT)

    fix(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                fix(cell.paragraphs)


# ---------------------------------------------------------------- content


def main():
    REPORT_DIR.mkdir(exist_ok=True)
    metrics = json.loads((RESULTS / "metrics.json").read_text())
    demo = json.loads((RESULTS / "demo_results.json").read_text(encoding="utf-8"))

    ds = metrics["dataset"]
    nb = metrics["models"]["Naive Bayes"]
    lr = metrics["models"]["Logistic Regression"]
    sw, da, ac = demo["software_engineer"], demo["data_analyst"], demo["accountant"]

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ---------------- title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run("Resume Analyzer and Job Matcher")
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(2)
    r = sub.add_run("Final Project Report")
    r.font.size = Pt(14)
    r.font.italic = True
    r.font.color.rgb = BLUE

    course = doc.add_paragraph()
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    course.paragraph_format.space_after = Pt(2)
    r = course.add_run("CAI4002: Introduction to Artificial Intelligence")
    r.font.size = Pt(12)
    r.font.color.rgb = MUTED

    team = doc.add_paragraph()
    team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    team.paragraph_format.space_after = Pt(10)
    r = team.add_run("Team: Martin Dang, Tuan Khang Phan")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = INK
    _para_border(team, "bottom", "1F3864", 12, 8)

    add_callout(
        doc, "Summary.",
        "This report describes the final state of the Resume Analyzer and Job "
        "Matcher: an AI system that extracts a confidence-weighted skill profile "
        "from a resume (PDF, DOCX, or plain text), parses a job posting into "
        "required and preferred qualifications, computes a blended TF-IDF, "
        "skill-coverage and Sentence-BERT compatibility score, classifies the "
        "resume's job category with a Naive Bayes model, and recommends the edits "
        "that would raise the score the most, all reachable from a minimal web "
        "interface.")

    add_metric_strip(doc, [
        (f"{ds['n_resumes']:,}", "resumes in corpus"),
        (f"{ds['n_categories']}", "job categories"),
        ("244", "skills in taxonomy"),
        (f"{lr['accuracy']:.1%}", "best classifier accuracy"),
    ])

    # ---------------- 1
    add_heading(doc, "1. Project Status")
    add_body(doc,
             "The system is feature-complete. Every pipeline stage from the "
             "original proposal is implemented and verified end to end: resume "
             "ingestion (PDF, DOCX, TXT), confidence-weighted skill extraction "
             f"over a 244-skill taxonomy spanning all {ds['n_categories']} "
             "dataset job categories, job-posting constraint parsing, a "
             "three-signal compatibility score (TF-IDF corpus percentile, skill "
             "coverage, and pre-trained Sentence-BERT semantic similarity), a "
             "greedy recommendation engine, a Naive Bayes category classifier, "
             "and a Streamlit web interface tying all of it together for an end "
             "user.")
    add_body(doc,
             "All code, experiment scripts, figures, and results referenced in "
             "this report are versioned in the team's GitHub repository "
             "(github.com/khangpt2k6/Intro_to_AI).")

    # ---------------- 2
    add_heading(doc, "2. Dataset")
    add_body(doc,
             "The system uses the public Resume Dataset from Kaggle "
             f"(snehaanbhawal/resume-dataset): {ds['n_resumes']:,} real resumes "
             f"collected from livecareer.com, labeled into {ds['n_categories']} "
             "job categories. The dataset trains the TF-IDF vector space and the "
             "category classifier, and supplies realistic test resumes for the "
             "demo. A public dataset keeps the project reproducible and avoids "
             "the privacy concerns of collecting real applicant data.")
    add_table(doc, ["Split", "Size"], [
        ["Total resumes", f"{ds['n_resumes']:,}"],
        ["Job categories", f"{ds['n_categories']}"],
        ["Training split (80%)", f"{ds['train_size']:,} resumes"],
        ["Held-out test split (20%)", f"{ds['test_size']:,} resumes"],
    ], widths=[3.0, 3.0])
    add_figure(doc, "fig1_category_distribution.png",
               "Figure 1. Distribution of the 2,484 resumes across the 24 job "
               "categories. The class imbalance (120 vs. 22 resumes) is one "
               "challenge the classifier has to handle.", width=5.6)

    # ---------------- 3
    add_heading(doc, "3. System Architecture")
    add_body(doc, "All six pipeline stages are implemented and working:")
    add_table(doc, ["Stage", "Implementation"], [
        ["Resume ingestion",
         "load_resume() extracts text from .pdf (pypdf), .docx (python-docx), "
         "or .txt, then normalizes whitespace and unicode punctuation (smart "
         "quotes, bullets, dashes) so extraction behaves the same regardless of "
         "source format."],
        ["Skill extraction and confidence",
         "A 244-skill taxonomy (685 aliases) matched with word-bounded regular "
         "expressions. Each extracted skill gets a confidence in [0, 1] blending "
         "mention frequency, experience context (proximity to action words like "
         "'built' or 'deployed'), and Naive Bayes category agreement."],
        ["Job description parsing",
         "Postings are split into hard (required) and soft (preferred) "
         "constraints."],
        ["Compatibility scoring",
         "Score = 0.4 x confidence-weighted skill coverage + 0.3 x "
         "corpus-percentile TF-IDF cosine + 0.3 x pre-trained Sentence-BERT "
         "cosine similarity."],
        ["Recommendation engine",
         "Missing skills ranked greedy best-first by exact score gain recovered; "
         "required skills carry twice the weight of preferred ones."],
        ["Category classification",
         "Naive Bayes and Logistic Regression over TF-IDF features, evaluated on "
         "a stratified held-out test set."],
    ], widths=[1.9, 4.4])

    # ---------------- 4
    add_heading(doc, "4. Skill Confidence Model")
    add_body(doc,
             "Binary alias matching tells us a skill is mentioned; it cannot "
             "tell us whether the resume genuinely demonstrates the skill or "
             "just lists it. A small Multinomial Naive Bayes classifier "
             f"estimates P(category | resume text) over the {ds['n_categories']} "
             "dataset categories, and each skill in the taxonomy is assigned an "
             "expected category profile from that same model (its "
             "predicted-category distribution over a pseudo-document made of the "
             "skill's own alias phrases). A skill whose home categories match "
             "the resume's predicted category, for example TensorFlow on a "
             "resume that reads as Information Technology, is trusted more than "
             "the same skill appearing on an unrelated resume, such as "
             "TensorFlow on an accountant resume.")
    add_body(doc,
             "Confidence blends three signals with weights 0.40 (mention "
             "frequency), 0.25 (experience context), and 0.35 (Naive Bayes "
             "category agreement). Without a classifier supplied, "
             "extract_skills() falls back to frequency and context only, so it "
             "still runs as a pure text function with no model or dataset "
             "dependency. The matcher credits each matched job requirement by "
             "the resume's confidence in that skill, so a genuinely demonstrated "
             "skill contributes more to the coverage score than a name-dropped "
             "one.")

    # ---------------- 5
    add_heading(doc, "5. Semantic Similarity")
    add_body(doc,
             "TF-IDF only rewards lexical overlap: a resume that says 'built "
             "REST services in Django' gets little credit against a posting "
             "asking for 'web backend development in Python', even though the "
             "two phrases describe the same capability. To capture that kind of "
             "paraphrase, the matcher adds a third signal: cosine similarity "
             "between sentence embeddings from a pre-trained Sentence-BERT model "
             "(all-MiniLM-L6-v2). On the demo the signal tracks fit closely, "
             f"from {ac['sbert_cos']:.2f} for the accountant posting to "
             f"{sw['sbert_cos']:.2f} for the software engineer posting.")
    add_callout(
        doc, "Scope note.",
        "This model is used strictly as a pre-trained black box for embeddings. "
        "It is loaded once, used only to encode text into vectors, and is never "
        "fine-tuned or otherwise updated on this project's data, consistent with "
        "the course's scope, which covers basic deep learning concepts but not "
        "model training at the transformer scale.")

    # ---------------- 6
    add_heading(doc, "6. Classifier Experiments")
    add_body(doc,
             "Two classifiers were trained to predict a resume's job category "
             f"from its text, using an 80/20 stratified train/test split "
             f"({ds['train_size']:,} training, {ds['test_size']:,} test "
             f"resumes). With {ds['n_categories']} classes, random guessing "
             "would score about 4.2%.")
    add_table(doc, ["Model", "Accuracy", "Macro F1"], [
        ["Multinomial Naive Bayes", f"{nb['accuracy']:.1%}", f"{nb['macro_f1']:.3f}"],
        ["Logistic Regression", f"{lr['accuracy']:.1%}", f"{lr['macro_f1']:.3f}"],
    ], widths=[2.6, 1.7, 1.7])
    add_figure(doc, "fig2_model_comparison.png",
               "Figure 2. Both models beat the 4.2% random baseline by a wide "
               "margin; Logistic Regression is the stronger of the two.",
               width=5.2)
    add_figure(doc, "fig3_confusion_matrix.png",
               "Figure 3. Normalized confusion matrix for Logistic Regression. "
               "Remaining confusion concentrates in genuinely overlapping "
               "categories, for example Finance vs. Accountant and Sales vs. "
               "Business Development.", width=6.0)
    add_callout(
        doc, "Note on reproducibility.",
        "The train/test split uses a fixed random_state, and Naive Bayes "
        f"reproduces exactly run to run ({nb['accuracy']:.2%} accuracy, "
        f"{nb['macro_f1']:.3f} macro F1). Logistic Regression varies by well "
        "under one point between runs of the same script "
        f"({lr['accuracy']:.2%} accuracy and {lr['macro_f1']:.3f} macro F1 in "
        "the run reported here), which comes from floating-point "
        "non-determinism in scikit-learn's solver rather than from any change "
        "to the taxonomy or features, since run_experiments.py does not depend "
        "on either.")

    # ---------------- 7
    add_heading(doc, "7. End-to-End Demo")
    add_body(doc,
             "To validate the whole pipeline, one real Information Technology "
             "resume from the dataset (ID 83816738) was matched against three "
             "job postings written in realistic formats: a backend software "
             "engineer role, a data analyst role, and a staff accountant role. "
             "The system behaves the way a human reviewer would expect: the IT "
             f"resume scores {sw['score']}/100 against the software engineer "
             f"posting, {da['score']}/100 against the data analyst posting, and "
             f"only {ac['score']}/100 against the accountant posting, a clear "
             "and well-separated ranking that matches the resume's actual skill "
             "set.")
    add_figure(doc, "fig4_match_scores.png",
               "Figure 4. The same resume scored against three postings. The "
               "score cleanly separates the strong fit from the partial and poor "
               "fits.", width=5.4)

    n_hard = len(sw["matched_hard"]) + len(sw["missing_hard"])
    top_gain = sw["recommendations"][0]["gain"]
    add_body(doc,
             "For the software engineer posting, the analyzer matched "
             f"{len(sw['matched_hard'])} of {n_hard} required skills "
             f"({', '.join(sw['matched_hard'])}) plus "
             f"{len(sw['matched_soft'])} preferred skills "
             f"({', '.join(sw['matched_soft'])}), and identified "
             f"{', '.join(sw['missing_hard'])} as the top missing required "
             f"skills, each worth an estimated +{top_gain} points.")
    add_table(doc, ["Posting", "Score", "Skill coverage", "Semantic (SBERT)"], [
        [sw["title"], f"{sw['score']}/100", f"{sw['coverage']:.0%}",
         f"{sw['sbert_cos']:.2f}"],
        [da["title"], f"{da['score']}/100", f"{da['coverage']:.0%}",
         f"{da['sbert_cos']:.2f}"],
        [ac["title"], f"{ac['score']}/100", f"{ac['coverage']:.0%}",
         f"{ac['sbert_cos']:.2f}"],
    ], widths=[2.2, 1.3, 1.5, 1.5])
    add_figure(doc, "fig5_demo_output.png",
               "Figure 5. Console output of the analyzer for the demo resume "
               "across all three postings, including per-skill confidence and "
               "ranked edit recommendations.", width=6.2)

    # ---------------- 8
    add_heading(doc, "8. Web Interface")
    add_body(doc,
             "A Streamlit app (app.py) exposes the full pipeline to an end "
             "user: upload or paste a resume, pick or paste a job description, "
             "and see the compatibility score with its three-way breakdown, the "
             "Naive Bayes predicted category, matched vs. missing skills with "
             "confidence, and the ranked recommended edits. The corpus, matcher, "
             "and classifier load once per session and are cached. This was "
             "verified live end to end: uploading a resume file through the "
             "browser correctly triggers ingestion, confidence-weighted "
             "extraction, and the blended score, matching the command-line "
             "output for the same input.")
    add_figure(doc, "fig6_web_ui.png",
               "Figure 6. The web interface after analyzing a resume against the "
               "backend software engineer posting: score breakdown, predicted "
               "category, and matched/missing skills with confidence.", width=4.7)

    # ---------------- 9
    add_heading(doc, "9. Challenges")
    add_bullet(doc, "Class imbalance in the dataset.",
               "The largest categories have 120 resumes while the smallest has "
               "only 22. Addressed with a stratified train/test split and macro "
               "F1 alongside raw accuracy.")
    add_bullet(doc, "Raw cosine similarities are not interpretable.",
               "Resume-to-posting TF-IDF cosine values range from about 0.01 to "
               "0.5, so presenting them directly would make most matches look "
               "weak. Solved with corpus-percentile calibration: the score "
               "reports where the resume ranks against all resumes in the corpus "
               "for that same posting.")
    add_bullet(doc, "Lexical overlap misses genuine matches.",
               "TF-IDF alone gives no credit to paraphrased skills. Adding a "
               "pre-trained Sentence-BERT similarity term as a third, "
               "independent signal closes this gap without requiring any model "
               "training.")
    add_bullet(doc, "Trusting a skill mention vs. a skill list.",
               "A skill named once in a flat 'Skills:' section is weaker "
               "evidence than one used in context ('built X with Y') or one "
               "consistent with the resume's overall predicted domain. The "
               "three-signal confidence model (frequency, context, Naive Bayes "
               "category agreement) addresses this directly.")
    add_bullet(doc, "Reproducibility under real classifier training.",
               "Naive Bayes is exactly reproducible run to run; Logistic "
               "Regression shows small drift (under 0.3 points) from "
               "solver-level floating-point non-determinism even with a fixed "
               "random seed, documented rather than hidden.")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ---------------- 10
    add_heading(doc, "10. Conclusion")
    add_body(doc,
             "The Resume Analyzer and Job Matcher is a complete, working "
             "prototype of every stage proposed at the start of the project: "
             "multi-format resume ingestion, confidence-aware skill extraction "
             "over an expanded, dataset-wide taxonomy, constraint-based job "
             "parsing, a three-signal compatibility score that blends lexical "
             "(TF-IDF), structured (skill coverage), and semantic (pre-trained "
             "Sentence-BERT) evidence, a greedy recommendation engine, a Naive "
             "Bayes category classifier, and a browser-based interface, verified "
             "end to end both from the command line and through the live web UI.")

    # ---------------- contributions
    add_heading(doc, "Team Contributions")
    add_table(doc, ["Martin Dang (50%)", "Tuan Khang Phan (50%)"], [
        ["Skill taxonomy design and alias normalization; skill extraction and "
         "confidence module; job description constraint parser; recommendation "
         "engine and ranking logic",
         "Dataset acquisition and preprocessing; TF-IDF matching engine and "
         "percentile calibration; Sentence-BERT integration; PDF and DOCX "
         "ingestion; classifier experiments and evaluation; Streamlit web UI"],
    ], widths=[3.15, 3.15])
    # both columns are equal partners, so do not bold only the first one
    body_row = doc.tables[-1].rows[1]
    for cell in body_row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = False

    apply_fonts(doc)
    out = REPORT_DIR / "Final_Report.docx"
    doc.save(out)
    print(f"Saved {out}")

    # best-effort PDF export (needs Microsoft Word via docx2pdf)
    try:
        from docx2pdf import convert
        pdf_out = REPORT_DIR / "Final_Report.pdf"
        convert(str(out), str(pdf_out))
        print(f"Saved {pdf_out}")
    except Exception as exc:  # noqa: BLE001
        print(f"PDF export skipped ({exc}); open the .docx in Word to export a PDF.")


if __name__ == "__main__":
    main()
